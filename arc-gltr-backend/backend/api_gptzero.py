import csv
import io
import json
import os
import re
import zipfile
from flask import send_file
import PyPDF2
import docx2txt
import docx
from docx.enum.text import WD_COLOR_INDEX
import requests
import zipfile


def extract_files(file):
    zip_files = []
    row=[['FileName','Status']]
    if file.filename.endswith('.docx') or file.filename.endswith('.pdf'):
        output_gpt = get_values(file, file.filename, zip_files)
        row.append(output_gpt)
    elif zipfile.is_zipfile(file):
        count_pdf_docx = 0
        with zipfile.ZipFile(file,'r') as zip:
            zip.extractall()
        for i in zip.infolist():
            if i.filename[0].isalpha() == True:
                if i.filename.endswith(".pdf") or i.filename.endswith(".docx") or i.filename.endswith(".txt") and 'MACOSX' not in i.filename:
                    count_pdf_docx += 1
                    output_gpt = get_values(i.filename, i.filename, zip_files)
                    row.append(output_gpt)
                if i.filename.endswith(".docx") == False:
                    os.remove('./'+i.filename)
        if count_pdf_docx == 0:
            print("No valid files in zip")
    else:
        print("Its not zip or pdf or docx or text file")
    
    output_file = "result.csv"
    with open(output_file, "w", newline="") as file:
        writer = csv.writer(file)
        writer.writerows(row)

    zip_files.append(output_file)
    in_memory_zip = io.BytesIO()
    with zipfile.ZipFile(in_memory_zip, 'w') as zip_file:
        for file in zip_files:
            zip_file.write(file)
            os.remove(file)

    in_memory_zip.seek(0)
    return send_file(in_memory_zip, download_name='result.zip', as_attachment=True)

def get_values(file, filename, zip_files):
    doc = docx.Document()
    para = doc.add_paragraph('''''')
    text = gettext(file, filename)
    file_name = [filename]
    filename_docx = f"{filename}"
    if filename.endswith('.docx'):
        filename_docx = filename[:len(filename_docx)-5]+'.docx'
    else:
        filename_docx = filename[:len(filename_docx)-4]+'.docx'
    status = check_gptzero(text,para) 
    if status == 0:
        file_name.append('Human Written')
    elif status < 0.5:
        file_name.append('Most likely Human Written')
    elif status < 1:
        file_name.append('Most likely AI/GPT Generated')
    elif status == 1:
        file_name.append('AI/GPT Generated')
 
    doc.save(filename_docx)    
    zip_files.append(filename_docx)
    return file_name       
 
def gettext(file, fileName):
    text = ""
    if fileName.endswith('.docx'):
        text = docx2txt.process(file)
    elif fileName.endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in range(len(reader.pages)):
            page_text = reader.pages[page]
            text += page_text.extract_text()
    return text

def check_gptzero(text, para):
    API_URL = 'https://api.gptzero.me/v2/predict/text'
    api_key = "01ba8f90e7b7475a9c164941bb3cbc0d"
    headers = {
    'X-API-KEY': api_key,
    'Content-Type': 'application/json'
    }
    data = {
        "document": text
    } 
    response = requests.post(API_URL, headers=headers, json=data)
    decoded_content = response.content.decode('utf-8')
    json_content = json.loads(decoded_content)
    for i in json_content['documents']:
        status = i['average_generated_prob']
        for j in range(len(i['sentences'])):
            if i['sentences'][j]['generated_prob'] == 1:
                para.add_run(i['sentences'][j]['sentence']+'. ').font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                para.add_run(i['sentences'][j]['sentence']+'. ')
    return status
